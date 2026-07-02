import streamlit as st
import fitz
import pytesseract
import pandas as pd
from PIL import Image
import io
import os
import re
import time
import traceback
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

from langchain_ollama import OllamaLLM
from langchain_core.prompts import PromptTemplate


# =====================================================================
# CONFIG
# =====================================================================
BASE_PATH = r"C:\AIFoundaryWorkshop"
OUTPUT_FOLDER = os.path.join(BASE_PATH, "Translated_Output")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Tesseract
TESS_PATH = os.path.join(BASE_PATH, "Tesseract-OCR", "tesseract.exe")
if os.path.exists(TESS_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESS_PATH

OCR_LANG = "jpn+eng"          # Japanese docs often contain numbers/codes in latin script
MIN_OCR_CONF = 30             # discard low-confidence OCR noise
DIGITAL_TEXT_THRESHOLD = 20   # chars; below this a PDF page is treated as "scanned"
COLUMN_CLUSTER_GAP = 40       # px gap used to decide a new table column


# =====================================================================
# UI THEME  (Opella Warm White - PANTONE 9285 C / #042B0B)
# =====================================================================
st.set_page_config(page_title="Opella AI Enterprise Translator", layout="wide")

st.markdown("""
<style>
/* Streamlit renders the app inside .stApp / stAppViewContainer, not <body> directly */
.stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
    background-color: #042B0B;
}
.block-container { padding: 2rem 3rem; }
.stButton>button {
    background-color: #042B0B;
    color: white;
    border-radius: 12px;
    font-weight: bold;
    border: none;
    padding: 0.6rem 1.2rem;
}
.stButton>button:hover {
    background-color: #0a4a19;
    color: white;
}
.card {
    background-color: white;
    padding: 20px;
    border-radius: 12px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08);
}
[data-testid="stSidebar"] {
    background-color: #F1E4D3;
}
</style>
""", unsafe_allow_html=True)

logo = os.path.join(BASE_PATH, "Opella_Logo.png")
if os.path.exists(logo):
    st.image(logo, width=180)


# =====================================================================
# TEXT CLEANUP
# =====================================================================
def clean_translation(text):
    if text is None:
        return ""
    text = str(text)
    text = re.sub(r'\*\*', '', text)
    text = re.sub(r'\*', '', text)
    text = re.sub(r'(?i)^\s*(translation|translated text)\s*:', '', text)
    text = text.strip()
    # strip wrapping quotes some models add
    if len(text) >= 2 and text[0] == text[-1] and text[0] in ('"', "'"):
        text = text[1:-1].strip()
    return text


def chunk_text(text, max_len=1200):
    """Split on line boundaries instead of raw character slicing so words
    are never cut in half mid-sentence."""
    lines = text.split("\n")
    chunks, current = [], ""
    for line in lines:
        if len(current) + len(line) + 1 > max_len and current:
            chunks.append(current)
            current = line
        else:
            current = f"{current}\n{line}" if current else line
    if current:
        chunks.append(current)
    return chunks if chunks else [""]


# =====================================================================
# TRANSLATION (Ollama)
# =====================================================================
_PROMPT = PromptTemplate.from_template("""
Translate the following Japanese text into clear, professional English.
Preserve structure and meaning. Return ONLY the translated text with no
preamble, no notes, and no quotation marks.

{text}
""")


def get_llm(model):
    return OllamaLLM(model=model)


def translate_raw(text, model, llm=None):
    """Translate a block of free text (used for .txt output)."""
    if not text or not text.strip():
        return ""
    llm = llm or get_llm(model)
    chain = _PROMPT | llm
    try:
        pieces = []
        for c in chunk_text(text):
            if not c.strip():
                continue
            pieces.append(chain.invoke({"text": c}))
        return clean_translation("\n".join(pieces))
    except Exception as e:
        st.warning(f"Translation error (falling back to original text): {e}")
        return text


def translate_cells(cells, model, llm=None, cache=None, progress_cb=None):
    """Translate a list of individual table cell strings, 1:1, so the
    output table always has exactly the same shape as the input table.
    A cache avoids re-translating repeated headers/units/values."""
    llm = llm or get_llm(model)
    cache = cache if cache is not None else {}
    chain = _PROMPT | llm
    out = []
    total = len(cells)
    for i, cell in enumerate(cells):
        stripped = str(cell).strip()
        if stripped == "":
            out.append("")
        elif stripped in cache:
            out.append(cache[stripped])
        else:
            try:
                result = clean_translation(chain.invoke({"text": stripped}))
            except Exception:
                result = stripped  # fail-safe: keep original rather than crash
            cache[stripped] = result
            out.append(result)
        if progress_cb:
            progress_cb((i + 1) / total)
    return out


# =====================================================================
# OCR + TABLE RECONSTRUCTION
# =====================================================================
def ocr_words(img):
    """Word-level OCR with tesseract's own line/block segmentation."""
    df = pytesseract.image_to_data(img, lang=OCR_LANG, output_type=pytesseract.Output.DATAFRAME)
    df["text"] = df["text"].astype(str)
    df = df[df["text"].str.strip() != ""]
    df["conf"] = pd.to_numeric(df["conf"], errors="coerce").fillna(-1)
    df = df[df["conf"] >= MIN_OCR_CONF]
    return df.reset_index(drop=True)


def _cluster_columns(lefts, gap=COLUMN_CLUSTER_GAP):
    """Cluster x-positions across the whole page into column centers."""
    uniq = sorted(set(int(x) for x in lefts))
    if not uniq:
        return []
    clusters, current = [], [uniq[0]]
    for x in uniq[1:]:
        if x - current[-1] <= gap:
            current.append(x)
        else:
            clusters.append(current)
            current = [x]
    clusters.append(current)
    return [sum(c) / len(c) for c in clusters]


def _assign_column(x, centers):
    return min(range(len(centers)), key=lambda i: abs(centers[i] - x))


def build_table_from_ocr(df):
    """Reconstruct a table from OCR word boxes using tesseract's line
    grouping for rows and x-position clustering for columns, instead of
    a fixed pixel-height guess."""
    if df.empty:
        return pd.DataFrame()

    df = df.copy()
    df["row_id"] = (
        df["block_num"].astype(str) + "_" +
        df["par_num"].astype(str) + "_" +
        df["line_num"].astype(str)
    )

    centers = _cluster_columns(df["left"].tolist())
    if not centers:
        return pd.DataFrame()

    row_order = df.groupby("row_id")["top"].min().sort_values().index.tolist()

    rows = []
    for row_id in row_order:
        group = df[df["row_id"] == row_id].sort_values("left")
        row_cells = {}
        for _, w in group.iterrows():
            col = _assign_column(w["left"], centers)
            row_cells[col] = (row_cells.get(col, "") + " " + w["text"]).strip() if col in row_cells else w["text"]
        rows.append([row_cells.get(c, "") for c in range(len(centers))])

    table = pd.DataFrame(rows)
    # drop fully-empty columns that can appear from stray noise clusters
    table = table.loc[:, (table != "").any(axis=0)]
    table.columns = range(table.shape[1])
    return table


# =====================================================================
# EXCEL FORMATTING
# =====================================================================
def format_excel(path):
    wb = load_workbook(path)
    ws = wb.active

    thin = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for row in ws.iter_rows():
        for cell in row:
            cell.border = thin
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    # bold header row with a solid fill (must use PatternFill, not a bare attribute)
    header_fill = PatternFill(start_color="042B0B", end_color="042B0B", fill_type="solid")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill

    # auto-fit column widths
    for col_cells in ws.columns:
        length = max((len(str(c.value)) if c.value is not None else 0) for c in col_cells)
        col_letter = get_column_letter(col_cells[0].column)
        ws.column_dimensions[col_letter].width = min(max(length + 4, 12), 60)

    ws.freeze_panes = "A2"
    wb.save(path)


def _finalize_table_columns(df):
    """Promote first row to header if it looks like one, name unnamed columns."""
    if df.empty:
        return df
    df = df.copy()
    header = df.iloc[0]
    body = df.iloc[1:].reset_index(drop=True)
    body.columns = [
        str(h).strip() if str(h).strip() else f"Column_{i+1}"
        for i, h in enumerate(header)
    ]
    return body


# =====================================================================
# FILE PROCESSORS
# =====================================================================
def _is_scanned_page(page):
    return len(page.get_text().strip()) < DIGITAL_TEXT_THRESHOLD


def process_pdf(path, model, llm, cache, status=None):
    doc = fitz.open(path)
    name = os.path.splitext(os.path.basename(path))[0]

    all_text_parts = []
    ocr_tables = []
    digital_lines = []
    any_scanned = False

    for page in doc:
        if _is_scanned_page(page):
            any_scanned = True
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img = Image.open(io.BytesIO(pix.tobytes()))
            df_words = ocr_words(img)
            if not df_words.empty:
                all_text_parts.append(" ".join(df_words["text"].tolist()))
                ocr_tables.append(build_table_from_ocr(df_words))
        else:
            page_text = page.get_text("text")
            all_text_parts.append(page_text)
            digital_lines.extend([ln for ln in page_text.split("\n") if ln.strip()])

    full_text = "\n".join(all_text_parts)

    if status:
        status.write("Translating extracted text...")
    translated_text = translate_raw(full_text, model, llm=llm)

    txt_path = os.path.join(OUTPUT_FOLDER, name + "_EN.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(translated_text)

    # Build the table: from OCR if the PDF was scanned, otherwise a
    # simple line-based table from the real text layer.
    if any_scanned and ocr_tables:
        raw_table = pd.concat([t for t in ocr_tables if not t.empty], ignore_index=True) if any(not t.empty for t in ocr_tables) else pd.DataFrame()
    else:
        raw_table = pd.DataFrame({0: ["Line"] + digital_lines}) if digital_lines else pd.DataFrame()

    excel_path = os.path.join(OUTPUT_FOLDER, name + "_EN.xlsx")
    if not raw_table.empty:
        df_final = _finalize_table_columns(raw_table)
        flat_cells = df_final.astype(str).values.flatten().tolist()
        if status:
            progress = status.progress(0.0)
            translated_cells = translate_cells(
                flat_cells, model, llm=llm, cache=cache,
                progress_cb=lambda p: progress.progress(p)
            )
        else:
            translated_cells = translate_cells(flat_cells, model, llm=llm, cache=cache)

        idx = 0
        for i in range(df_final.shape[0]):
            for j in range(df_final.shape[1]):
                df_final.iat[i, j] = translated_cells[idx]
                idx += 1

        df_final.to_excel(excel_path, index=False)
        format_excel(excel_path)
    else:
        # still produce a valid (empty-but-structured) workbook rather than nothing
        pd.DataFrame({"Text": []}).to_excel(excel_path, index=False)
        format_excel(excel_path)

    return txt_path, excel_path


def process_docx(path, model, llm, cache, status=None):
    doc = Document(path)
    name = os.path.splitext(os.path.basename(path))[0]

    # --- text (paragraphs) ---
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if status:
        status.write("Translating document text...")
    translated_text = translate_raw(text, model, llm=llm)

    txt_path = os.path.join(OUTPUT_FOLDER, name + "_EN.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(translated_text)

    # --- tables (use real Word tables if present, else fall back to
    # a simple one-column table built from paragraphs) ---
    excel_path = os.path.join(OUTPUT_FOLDER, name + "_EN.xlsx")

    if doc.tables:
        table = doc.tables[0]
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        raw_table = pd.DataFrame(rows)
    elif text:
        raw_table = pd.DataFrame({0: ["Line"] + [ln for ln in text.split("\n") if ln.strip()]})
    else:
        raw_table = pd.DataFrame()

    if not raw_table.empty:
        df_final = _finalize_table_columns(raw_table)
        flat_cells = df_final.astype(str).values.flatten().tolist()
        translated_cells = translate_cells(flat_cells, model, llm=llm, cache=cache)
        idx = 0
        for i in range(df_final.shape[0]):
            for j in range(df_final.shape[1]):
                df_final.iat[i, j] = translated_cells[idx]
                idx += 1
        df_final.to_excel(excel_path, index=False)
        format_excel(excel_path)

    return txt_path, (excel_path if not raw_table.empty else None)


def process_excel(path, model, llm, cache, status=None):
    name = os.path.splitext(os.path.basename(path))[0]
    xls = pd.read_excel(path, sheet_name=None)

    excel_path = os.path.join(OUTPUT_FOLDER, name + "_EN.xlsx")
    all_text_parts = []

    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        for sheet, df in xls.items():
            df = df.fillna("").astype(str)
            all_text_parts.append("\n".join(df.values.flatten().tolist()))

            flat = df.values.flatten().tolist()
            if status:
                status.write(f"Translating sheet '{sheet}'...")
            translated = translate_cells(flat, model, llm=llm, cache=cache)

            idx = 0
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    df.iat[i, j] = translated[idx]
                    idx += 1

            df.to_excel(writer, sheet_name=sheet, index=False)

    format_excel(excel_path)

    txt_path = os.path.join(OUTPUT_FOLDER, name + "_EN.txt")
    translated_text = translate_raw("\n".join(all_text_parts), model, llm=llm)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(translated_text)

    return txt_path, excel_path


# =====================================================================
# FOLDER PICKER (optional, works only when run locally with a display)
# =====================================================================
def _browse_folder():
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        folder = filedialog.askdirectory()
        root.destroy()
        return folder
    except Exception:
        return None


# =====================================================================
# MAIN UI
# =====================================================================
def main():
    st.title("🚀 Opella AI Enterprise Translator")
    st.caption("Japanese → English document translation with table extraction")

    col1, col2 = st.columns([3, 1])
    with col1:
        model = st.selectbox("Ollama Model", ["llama2", "llama3", "mistral"], index=0)
    with col2:
        st.write("")

    if "folder_path" not in st.session_state:
        st.session_state.folder_path = BASE_PATH

    fcol1, fcol2 = st.columns([4, 1])
    with fcol1:
        folder_path = st.text_input("📂 Folder Path", st.session_state.folder_path)
    with fcol2:
        st.write("")
        if st.button("Browse..."):
            picked = _browse_folder()
            if picked:
                st.session_state.folder_path = picked
                st.rerun()
            else:
                st.info("Folder browser isn't available in this environment — type the path directly.")

    st.markdown(f"Output will be saved to: `{OUTPUT_FOLDER}`")

    if st.button("Process All Files 🚀"):
        if not os.path.isdir(folder_path):
            st.error(f"Folder not found: {folder_path}")
            return

        files = [f for f in os.listdir(folder_path)
                  if os.path.isfile(os.path.join(folder_path, f))
                  and f.lower().endswith((".pdf", ".docx", ".xlsx"))]

        if not files:
            st.warning("No .pdf, .docx, or .xlsx files found in that folder.")
            return

        start = time.time()
        llm = get_llm(model)
        cache = {}  # shared translation cache across all files in this run
        results = []

        overall_progress = st.progress(0.0)

        for i, f in enumerate(files):
            full = os.path.join(folder_path, f)
            with st.status(f"Processing {f}...", expanded=False) as status:
                try:
                    if f.lower().endswith(".pdf"):
                        txt_out, xlsx_out = process_pdf(full, model, llm, cache, status=status)
                    elif f.lower().endswith(".docx"):
                        txt_out, xlsx_out = process_docx(full, model, llm, cache, status=status)
                    elif f.lower().endswith(".xlsx"):
                        txt_out, xlsx_out = process_excel(full, model, llm, cache, status=status)
                    else:
                        continue
                    results.append((f, txt_out, xlsx_out, None))
                    status.update(label=f"✅ {f} done", state="complete")
                except Exception as e:
                    err = f"{e}\n{traceback.format_exc(limit=2)}"
                    results.append((f, None, None, str(e)))
                    status.update(label=f"❌ {f} failed", state="error")
                    st.error(f"Error processing {f}: {e}")

            overall_progress.progress((i + 1) / len(files))

        st.success(f"✅ Completed {len(files)} file(s) in {round(time.time() - start, 2)} sec")

        st.subheader("Results")
        for f, txt_out, xlsx_out, err in results:
            with st.container():
                st.markdown(f"**{f}**")
                if err:
                    st.error(err)
                    continue
                dcols = st.columns(2)
                if txt_out and os.path.exists(txt_out):
                    with open(txt_out, "rb") as fh:
                        dcols[0].download_button(
                            "⬇️ Download Text", fh.read(),
                            file_name=os.path.basename(txt_out), key=f"txt_{f}"
                        )
                if xlsx_out and os.path.exists(xlsx_out):
                    with open(xlsx_out, "rb") as fh:
                        dcols[1].download_button(
                            "⬇️ Download Excel", fh.read(),
                            file_name=os.path.basename(xlsx_out), key=f"xlsx_{f}"
                        )


if __name__ == "__main__":
    main()